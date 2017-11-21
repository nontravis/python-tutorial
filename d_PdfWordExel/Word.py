import docx

#┌───────────────┐
#│ NOTE: Reading │
#└───────────────┘
d = docx.Document("d_PdfWordExel/demo.docx")
d.paragraphs
for pnum in range(len(d.paragraphs)):
    d.paragraphs[pnum].text

p = d.paragraphs[1]
p.runs      # seperate by change style
for lnum in range(len(p.runs)):
    p.runs[lnum].text

p.runs[0].bold == None
p.runs[1].bold
p.runs[3].italic

#┌────────────────┐
#│ NOTE: Editting │
#└────────────────┘
p.runs[3].underline = True
p.runs[3].text = "italic and underlined."
d.save("d_PdfWordExel/demo2.docx")

p.style = "Title"
d.save("d_PdfWordExel/demo2.docx")

d = docx.Document()
d.add_paragraph("Hello this is a paragraph.")
d.add_paragraph("This is another paragraph.")
d.save("d_PdfWordExel/demo3.docx")
p = d.paragraphs[0]
p.add_run("This is new run.")
p.runs[1].bold = True
d.save("d_PdfWordExel/demo3.docx")
